import io
import os
import itertools
import copy

from utils.OrderedSet import OrderedSet
from utils.dataObjects import DocxWordData
from utils.validators import (
    file_validate,
    number_validate
)
from utils.constants import (
    russianAlphabet,
    englishAlphbet,
    digits,
)
from utils.isolate_run import isolate_run
from find_subtext import compareTwoTexts

from tkinter import *
import tkinter as tk
from tkinter import filedialog as fd
import docxpy
import fitz
from docx import Document
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX


def extract_words_from_docx(document: Document, alphabet) -> list[DocxWordData]:
    arr = []
    temp = []
    curr_paragraph = 0
    curr_symbol_pos = 0
    word_number = 0
    for paragraph in document.paragraphs:
        for char in paragraph.text.lower():
            if char in alphabet or char == '-' and len(temp) > 0:
                temp.append(char)
                curr_symbol_pos += 1
            else:
                if len(temp) > 0:
                    word = DocxWordData(
                        word=''.join(temp),
                        paragraph_num=curr_paragraph,
                        first_symbol_pos=curr_symbol_pos - len(temp),
                        wnumber=word_number,
                    )
                    arr.append(word)
                    word_number += 1
                    curr_symbol_pos += 1
                    temp = []
        if len(temp) > 0:
            word = DocxWordData(
                word=''.join(temp),
                paragraph_num=curr_paragraph,
                first_symbol_pos=curr_symbol_pos - len(temp),
                wnumber=word_number,
            )
            arr.append(word)
        curr_paragraph += 1
        curr_symbol_pos = 0
        temp = []
    return arr


def color_subtext(run: Run):
    fontt = run.font
    fontt.highlight_color = WD_COLOR_INDEX.YELLOW


def make_runs_from_words(document: Document, start_word: DocxWordData, end_word: DocxWordData) -> list[Run]:
    runs = []
    if (start_word.paragraph_num == end_word.paragraph_num) and (start_word.wnumber < end_word.wnumber):
        run = isolate_run(
            document.paragraphs[start_word.paragraph_num],
            start_word.first_symbol_pos,
            end_word.first_symbol_pos + len(end_word.word),
        )
        runs.append(run)
        return runs
    elif start_word.paragraph_num < end_word.paragraph_num:
        first_run = isolate_run(
            document.paragraphs[start_word.paragraph_num],
            start_word.first_symbol_pos,
            len(document.paragraphs[start_word.paragraph_num].text)
        )
        last_run = isolate_run(
            document.paragraphs[end_word.paragraph_num],
            0,
            end_word.first_symbol_pos + len(end_word.word)
        )
        if end_word.paragraph_num - start_word.paragraph_num == 1:
            runs.append(first_run)
            runs.append(last_run)
            return runs
        else:
            runs.append(first_run)
            for i in range(start_word.paragraph_num + 1, end_word.paragraph_num):
                if len(document.paragraphs[i].text) > 0:
                    cur_run = isolate_run(
                        document.paragraphs[i],
                        0,
                        len(document.paragraphs[i].text)
                    )
                    runs.append(cur_run)
            runs.append(last_run)
            return runs
    else:
        raise ValueError('Последнее слово находится до первого')


def docx_words_to_txt(docx_words: list[DocxWordData]):
    words = []
    for word in docx_words:
        words.append(word.word)
    return words


def compare_two_docx(docx_words1: list[DocxWordData], docx_words2: list[DocxWordData], alphabet):
    txt1 = ' '.join(docx_words_to_txt(docx_words1))
    txt2 = ' '.join(docx_words_to_txt(docx_words2))
    rez = compareTwoTexts(txt1, txt2, alphabet)
    return rez


def main():
    alphabet = set.union(russianAlphabet, englishAlphbet, digits)
    min_words = 15
    path_file1 = 'I:/develop/ОТЧЕТ_ИТОГ_01_12_2021_в_отправку_в_Минэнерго.docx'
    path_file2 = 'I:/develop/ОТЧЕТ_ЭНЕРДЖИНЕТ_ФИНАЛ_в_отправку_13_12_2022.docx'
    # path_file1 = 'I:/develop/test2.docx'
    # path_file2 = 'I:/develop/test1.docx'
    f = open(path_file1, 'rb')
    document1 = Document(f)
    f.close()
    f = open(path_file2, 'rb')
    document2 = Document(f)
    f.close()
    docx_first_last_words = []
    words1 = extract_words_from_docx(document1, alphabet)
    words2 = extract_words_from_docx(document2, alphabet)
    if len(words1) <= len(words2):
        document = document1
        words = words1
        subtexts = compare_two_docx(words, words2, alphabet)
    else:
        document = document2
        words = words2
        subtexts = compare_two_docx(words, words1, alphabet)
    for subtext in subtexts:
        if len(subtext) >= min_words - 1:
            docx_first_last_words.append((subtext[0][0], subtext[-1][0]))
    for sub in docx_first_last_words:
        runs = make_runs_from_words(document, words[sub[0]], words[sub[1] + 1])
        for run in runs:
            color_subtext(run)
    document.save('rezult.docx')


if __name__ == '__main__':
    main()
